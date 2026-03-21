#!/usr/bin/env python3
"""
備用 Markdown → Word 轉換腳本
在 markdown-to-word-template 技能不可用時使用

支援內建模板：tech-spec, formal-doc, user-guide, minimal
支援變數替換：{{project_name}}, {{version}}, {{author}}, {{date}}, {{status}}, {{company}}
"""

import sys
import re
import argparse
from datetime import date

def get_builtin_theme(name):
    themes = {
        "tech-spec": {
            "h1_color": "1F4E79", "h1_size": 32,
            "h2_color": "2E74B5", "h2_size": 28,
            "h3_color": "2E74B5", "h3_size": 24,
            "font": "Segoe UI",
            "table_header_color": "1F4E79",
        },
        "formal-doc": {
            "h1_color": "2E2E2E", "h1_size": 32,
            "h2_color": "404040", "h2_size": 28,
            "h3_color": "595959", "h3_size": 24,
            "font": "Arial",
            "table_header_color": "2E2E2E",
        },
        "user-guide": {
            "h1_color": "1E5C1E", "h1_size": 32,
            "h2_color": "2D862D", "h2_size": 28,
            "h3_color": "3D9B3D", "h3_size": 24,
            "font": "Calibri",
            "table_header_color": "1E5C1E",
        },
        "minimal": {
            "h1_color": "000000", "h1_size": 32,
            "h2_color": "000000", "h2_size": 28,
            "h3_color": "000000", "h3_size": 24,
            "font": "Arial",
            "table_header_color": "333333",
        },
    }
    return themes.get(name, themes["tech-spec"])

def replace_variables(text, variables):
    for key, value in variables.items():
        text = text.replace(f"{{{{{key}}}}}", value)
    return text

def parse_frontmatter(text):
    """Extract YAML frontmatter variables"""
    variables = {}
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            fm = text[3:end].strip()
            text = text[end+3:].strip()
            for line in fm.split("\n"):
                if ":" in line:
                    k, _, v = line.partition(":")
                    variables[k.strip()] = v.strip().strip('"')
    return text, variables

def convert_md_to_docx(md_path, output_path, builtin=None, template=None):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree
    except ImportError:
        print("Installing python-docx...")
        import subprocess
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches

    theme = get_builtin_theme(builtin or "tech-spec")

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Parse frontmatter
    content, fm_vars = parse_frontmatter(content)

    # Default variables
    variables = {
        "date": str(date.today()),
        "version": "1.0",
        "status": "Draft",
        "company": "",
        "author": "",
        "project_name": "",
        **fm_vars
    }
    content = replace_variables(content, variables)

    # Create document
    if template:
        doc = Document(template)
        # Clear body
        for element in doc.element.body[:]:
            if element.tag != qn('w:sectPr'):
                doc.element.body.remove(element)
    else:
        doc = Document()

    # Set default font
    doc.styles['Normal'].font.name = theme["font"]
    doc.styles['Normal'].font.size = Pt(11)

    # Set heading styles
    def set_heading_style(style_name, size, color_hex):
        try:
            style = doc.styles[style_name]
            style.font.name = theme["font"]
            style.font.size = Pt(size / 2)  # Pt conversion
            r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
            style.font.color.rgb = RGBColor(r, g, b)
            style.font.bold = True
        except Exception:
            pass

    set_heading_style('Heading 1', theme["h1_size"], theme["h1_color"])
    set_heading_style('Heading 2', theme["h2_size"], theme["h2_color"])
    set_heading_style('Heading 3', theme["h3_size"], theme["h3_color"])

    def hex_to_rgb(hex_color):
        return int(hex_color[0:2],16), int(hex_color[2:4],16), int(hex_color[4:6],16)

    def add_run_with_format(para, text, bold=False, italic=False, code=False):
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        return run

    def process_inline(para, text):
        """Process inline formatting: bold, italic, code"""
        # Split by inline code first
        parts = re.split(r'(`[^`]+`)', text)
        for part in parts:
            if part.startswith('`') and part.endswith('`'):
                add_run_with_format(para, part[1:-1], code=True)
            else:
                # Process bold/italic
                segments = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__)', part)
                for seg in segments:
                    if seg.startswith('**') and seg.endswith('**'):
                        add_run_with_format(para, seg[2:-2], bold=True)
                    elif seg.startswith('*') and seg.endswith('*'):
                        add_run_with_format(para, seg[1:-1], italic=True)
                    else:
                        add_run_with_format(para, seg)

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('# '):
            para = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith('## '):
            para = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith('### '):
            para = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        elif line.startswith('#### '):
            para = doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        # Horizontal rule
        elif re.match(r'^[-*]{3,}$', line.strip()):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), theme["h2_color"])
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote
        elif line.startswith('> '):
            para = doc.add_paragraph(style='Quote' if 'Quote' in [s.name for s in doc.styles] else 'Normal')
            process_inline(para, line[2:])
            para.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue

        # Code block
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            para = doc.add_paragraph()
            run = para.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.3)
            # Light gray background via shading
            try:
                pPr = para._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')
                pPr.append(shd)
            except Exception:
                pass
            continue

        # Table
        elif '|' in line and i + 1 < len(lines) and re.match(r'^[|\s\-:]+$', lines[i+1]):
            # Parse table
            table_lines = [line]
            i += 1  # skip separator
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1

            if not table_lines:
                continue

            # Parse header
            headers = [c.strip() for c in table_lines[0].split('|') if c.strip()]
            if not headers:
                continue

            t = doc.add_table(rows=1, cols=len(headers))
            t.style = 'Table Grid'

            # Header row
            hdr_cells = t.rows[0].cells
            for j, h in enumerate(headers):
                hdr_cells[j].text = h
                # Style header
                for run in hdr_cells[j].paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = theme["font"]
                # Set header background
                try:
                    tc = hdr_cells[j]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), theme["table_header_color"])
                    tcPr.append(shd)
                except Exception:
                    pass

            # Data rows
            for row_idx, row_line in enumerate(table_lines[1:]):
                cells = [c.strip() for c in row_line.split('|') if c.strip()]
                row = t.add_row()
                for j, cell_text in enumerate(cells):
                    if j < len(headers):
                        row.cells[j].text = cell_text
                        # Alternating row color
                        if row_idx % 2 == 1:
                            try:
                                tc = row.cells[j]._tc
                                tcPr = tc.get_or_add_tcPr()
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), 'F0F0F0')
                                tcPr.append(shd)
                            except Exception:
                                pass

            doc.add_paragraph()  # spacing after table
            continue

        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            process_inline(para, line[2:])
            i += 1
            continue

        # Numbered list
        elif re.match(r'^\d+\. ', line):
            para = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line)
            process_inline(para, text)
            i += 1
            continue

        # Checkbox list
        elif line.startswith('- [ ]') or line.startswith('- [x]') or line.startswith('- [X]'):
            checked = '[x]' in line.lower()
            text = line[6:].strip()
            para = doc.add_paragraph(style='List Bullet')
            para.add_run('☑ ' if checked else '☐ ')
            process_inline(para, text)
            i += 1
            continue

        # Empty line
        elif line.strip() == '':
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        else:
            if line.strip():
                para = doc.add_paragraph()
                process_inline(para, line)
            i += 1
            continue

    doc.save(output_path)
    print(f"Saved: {output_path}")

def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to Word (.docx)')
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('-o', '--output', required=True, help='Output .docx file path')
    parser.add_argument('--builtin', choices=['tech-spec', 'formal-doc', 'user-guide', 'minimal'],
                        default='tech-spec', help='Built-in theme')
    parser.add_argument('--template', help='Path to .docx template file')
    args = parser.parse_args()

    convert_md_to_docx(args.input, args.output, builtin=args.builtin, template=args.template)

if __name__ == '__main__':
    main()
